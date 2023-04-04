# -*- coding: utf-8 -*-
# 农事服务 春播气象服务
import os
import sys
import json
import datetime
import psycopg2
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from base import BaseProduct

text_type = str
if sys.version[:1] < '3':
    text_type = unicode
    reload(sys)
    sys.setdefaultencoding('utf8')


class Doc141(BaseProduct):
    def __init__(self, *args, **kwargs):
        BaseProduct.__init__(self, *args, **kwargs)
        self.red_header = kwargs.get("red_header")

    @staticmethod
    def set_table_label(cell, label):
        """
        设置表格表头
        :param cell: 单元格对象
        :param label: str, 表头内容
        :return:
        """
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label)
        run.font.size = Pt(12)  # 小四
        run.font.name = "仿宋"
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中

    @staticmethod
    def get_win_level(win_speed_min, win_speed_max):
        """
        根据风速判断风级
        :param win_speed_min: float, 最小风速 m/s
        :param win_speed_max: float, 最大风速 m/s
        :return: win_level_min 最小风级, win_level_max 最大风级
        """
        win_level_dict = {0: [0, 0.2], 1: [0.3, 1.5], 2: [1.6, 3.3], 3: [3.4, 5.4], 4: [5.5, 7.9],
                          5: [8.0, 10.7], 6: [10.8, 13.8], 7: [13.9, 17.1], 8: [17.2, 20.7], 9: [20.8, 24.4],
                          10: [24.5, 28.4], 11: [28.5, 32.6], 12: [37.0, 41.4], }
        win_level_min, win_level_max = None, None
        for k, v in win_level_dict.items():
            if v[0] <= win_speed_min <= v[1]:
                win_level_min = k
            if v[0] <= win_speed_max <= v[1]:
                win_level_max = k
        return win_level_min, win_level_max

    @staticmethod
    def get_win_direction(win_d):
        """
        根据风向角度判断风向
        :param win_d: float, 风向角度
        :return: win_d_str: str, 风向
        """
        win_direction_dict = {
            "东北风": [11.26, 78.75],
            "东风": [78.76, 101.25],
            "东南风": [101.26, 168.75],
            "南风": [168.76, 191.25],
            "西南风": [191.26, 258.75],
            "西风": [258.76, 281.25],
            "西北风": [281.26, 348.76],
        }
        for k, v in win_direction_dict.items():
            if win_d >= 348.76 and win_d <= 11.25:
                return "北风"
            elif v[0] <= win_d and win_d <= v[1]:
                return k

    @staticmethod
    def set_table_cell(cell, label):
        """
        设置表格内容
        :param cell: 单元格对象
        :param label: str, 表格内容
        :return:
        """
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label)
        run.font.size = Pt(10)  # 小四
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中

    def load_style(self, document):
        super(Doc141, self).load_style(document)

        document.styles.add_style("New Header", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["New Header"].font.name = u"方正小标宋简体"
        document.styles["New Header"].element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
        document.styles["New Header"].font.size = Pt(28)
        document.styles["New Header"].font.color.rgb = RGBColor(255, 0, 0)  # 红色
        document.styles["New Header"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.styles["New Header"].paragraph_format.space_after = Pt(0)  # 段落间距
        document.styles["New Header"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 单倍行距

        document.styles.add_style("New SignerLine", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["New SignerLine"].font.name = u'楷体_GB2312'
        document.styles["New SignerLine"].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体_GB2312')
        document.styles["New SignerLine"].font.size = Pt(14)  # 四号
        document.styles["New SignerLine"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.styles["New SignerLine"].paragraph_format.space_after = Pt(0)  # 段落间距
        document.styles["New SignerLine"].paragraph_format.line_spacing = Pt(35)  # 固定行距

        document.styles.add_style("New Footer", WD_STYLE_TYPE.CHARACTER)
        document.styles["New Footer"].font.name = u'楷体_GB2312'
        document.styles["New Footer"].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体_GB2312')
        document.styles["New Footer"].font.size = Pt(14)  # 四号

        document.styles.add_style("New Heading P", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["New Heading P"].font.name = u'黑体'
        document.styles["New Heading P"].element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.styles["New Heading P"].font.size = Pt(16)  # 三号
        document.styles["New Heading P"].paragraph_format.space_after = Pt(0)  # 段落间距
        document.styles["New Heading P"].paragraph_format.line_spacing = Pt(28)  # 固定行距

        document.styles.add_style("Main Text", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["Main Text"].font.name = u'仿宋'
        document.styles["Main Text"].element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        document.styles["Main Text"].font.size = Pt(16)  # 三号
        document.styles["Main Text"].paragraph_format.space_after = Pt(16)  # 段落间距
        document.styles["Main Text"].paragraph_format.line_spacing = Pt(28)  # 固定行距
        document.styles["Main Text"].paragraph_format.first_line_indent = Pt(32)  # 首行缩进

    def make_doc(self):
        document = Document()
        self.load_style(document)

        self.set_footer(document)
        self.set_reference(document)
        self.insert_top_pic(document, 15.09, 1.88)
        title = document.add_paragraph(u"内蒙古生态与农业气象信息", style="Body Text 3")
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ISSN = document.add_paragraph(u"农业气象—第{}期  （总{}期）  ".format(self.issue, self.issue_total),
                                      style="Body Text 2")
        ISSN.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.set_sign_date(document)
        self.insert_red_line_pic(document, 16.56, 0.09)

        label = document.add_paragraph(text_type(self.label), style="label")
        label.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        abstract = document.add_paragraph(content_head, style='Abstract')
        abstract.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(u"    一、春播农业气象条件", style="Heading P")
        document.add_paragraph(content_p1)

        self.insert_pic(document, u"./per_tmp.jpg", 11.35, 7.96)
        desc_pic1 = document.add_paragraph(u"图1  3月11日至5月31日内蒙古平均气温分布图（℃）", style="pic_desc")
        desc_pic1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.insert_pic(document, u"./tmp_departure.jpg", 11.35, 7.96)
        desc_pic2 = document.add_paragraph(u"图2  3月11日至5月31日内蒙古温度距平图（℃）", style="pic_desc")
        desc_pic2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.insert_pic(document, u"./min_tmp.jpg", 11.35, 7.96)
        desc_pic3 = document.add_paragraph(u"图3   3月11日至5月31日内蒙古降水量图(mm）", style="pic_desc")
        desc_pic3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.insert_pic(document, u"./rain_fall.jpg", 11.35, 7.96)
        desc_pic4 = document.add_paragraph(u"图4   3月11日至5月31日内蒙古降水距平百分率图(%）", style="pic_desc")
        desc_pic4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.insert_pic(document, u"./rain_departure.jpg", 11.35, 7.96)
        desc_pic5 = document.add_paragraph(u"图5内蒙古地区3月中旬至5月下旬大于10度有效积温（℃·d）", style="pic_desc")
        desc_pic5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_paragraph(u"    二、农业生产形势", style="Heading P")
        document.add_paragraph(content_p2)

        self.insert_pic(document, u'./soil_moisture.jpg', 10.55, 9.3)
        desc_pic8 = document.add_paragraph(u"图6  5月下旬内蒙古农区土壤墒情分布图", style="pic_desc")
        desc_pic8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_paragraph(u"")
        desc_table1 = document.add_paragraph(u"表  2015年5月下旬农区土壤墒情面积百分比  单位：%", style="pic_desc")
        desc_table1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_table(9, 13)

        document.add_paragraph(u"    三、播种进度及作物发育状况", style="Heading P")
        document.add_paragraph(content_p3)

        self.insert_pic(document, u'./春小麦发育期.jpg', 12.33, 10.9)
        desc_pic8 = document.add_paragraph(u"图7  春小麦发育期", style="pic_desc")
        desc_pic8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.insert_pic(document, u"./春玉米发育期.jpg", 12.49, 11.03)
        desc_pic8 = document.add_paragraph(u"图8  春玉米发育期", style="pic_desc")
        desc_pic8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_paragraph(u"    四、农业气象条件展望与建议", style="Heading P")
        document.add_paragraph(content_p4)

        document.save('./{}.docx'.format(self.uuid))

    def make_doc_for_country(self):
        document = Document()
        self.load_style(document)
        # 页脚
        footer = document.sections[0].footer.paragraphs[0]
        footer.add_run().add_picture(u"./thin_red_line.jpg", width=Cm(16.56), height=Cm(0.03))
        text = f"值班预报员：{self.main_class}                        签发人：{self.signer}"
        footer.add_run(text, style="New Footer")
        # 红头
        document.add_paragraph(f"{self.red_header}", style="New Header")
        # 单位、时间
        dt = f"{self.date[0]}年{int(self.date[1])}月{int(self.date[2])}日"
        text = f"{self.unit}                                {dt}"
        document.add_paragraph(text, style="New SignerLine")
        self.insert_red_line_pic(document, 16.56, 0.09)

        # 查询数据
        start_time = datetime.datetime.strptime(f"{self.date[0]}{self.date[1]}{self.date[2]}", "%Y%m%d")
        start_time_str = start_time.strftime("%Y-%m-%d %H:%M:%S")
        station_list = "','".join(list(self.station_dict.keys()))
        valid_time_list = "','".join([str(i * 24) for i in range(1, 8)])
        sql = f"SELECT tem, win_s, win_d, wep, validtime FROM sevp_wefc " \
              f"WHERE datetime = '{start_time_str}' " \
              f"and station_id_c in ('{station_list}') and validtime in ('{valid_time_list}')"
        conn = psycopg2.connect(self.dsn)
        with conn.cursor() as cur:
            cur.execute(sql)
            columns = [str(row[0]) for row in cur.description]
            # 组装dataFrame
            weather_info = pd.DataFrame(cur.fetchall(), columns=columns)
            for column in weather_info.columns[:-1]:
                # decimal 转 float
                weather_info[column] = weather_info[column].apply(float)
            weather_info[weather_info == 999.9] = np.nan
            # 以有效时间聚合保留极值
            weather_info = weather_info.groupby("validtime").agg({
                "tem": ["max", "min"],  # 最高温度 最低温度
                "win_s": ["max", "min"],  # 最大风速 最小风速
                "win_d": ["max", "min"],  # 风向范围
                "wep": ["max", "min"]  # 天气变化
            })
            weather_info.reset_index(inplace=True)
            date_list = []
            for _ in weather_info["validtime"]:
                date_list.append((start_time + datetime.timedelta(days=int(_ // 24))).day)
            weather_info["validtime"] = date_list
            weather_info.columns = ["date", "tem_max", "tem_min", "win_s_max", "win_s_min", "win_d_max", "win_d_min", "wep_max", "wep_min"]
            weather_info.to_csv("weather_info.csv")

        # 首段描述
        document.add_paragraph("    一、未来七天天气趋势", style="New Heading P")
        # 文字说明
        date_weather_info = weather_info.set_index("date")
        text = ""
        label_name_dict = {"tem": "温度", "win_s": "风速"}
        label_unit_dict = {"tem": "℃", "win_s": "m/s"}
        for k, v in label_name_dict.items():
            max_v = date_weather_info[f"{k}_max"].max()
            max_i = date_weather_info[f"{k}_max"].idxmax()
            min_v = date_weather_info[f"{k}_min"].min()
            min_i = date_weather_info[f"{k}_min"].idxmin()
            unit = label_unit_dict[k]
            text += f"最高{v}出现在{max_i}日，为{max_v}{unit}，最低{v}出现在{min_i}日，为{min_v}{unit}。"
        document.add_paragraph(text, style="Main Text")

        # 天气预报
        paragraph = document.add_paragraph("    二、未来七天具体预报", style="New Heading P")
        paragraph.paragraph_format.space_after = Pt(16)  # 段后1行
        # 添加表格
        table = document.add_table(9, 5, style="Table Grid")
        # 设置表格样式
        for row in table.rows:
            row.height = Cm(1)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 合并单元格
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(1, 1))
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 2).merge(table.cell(0, 3))
        # 填充表头
        col_width_dict = {0: 1.79, 1: 3.68, 2: 3.68, 3: 3.68, 4: 3.06}  # 单元格宽度
        labels = ["日期", "天气", self.region['name'], "", "风向"]
        for i in range(len(labels)):
            cell = table.cell(0, i)
            self.set_table_label(cell, labels[i])
            cell.width = Cm(col_width_dict[i])  # 设定单元格宽度
        self.set_table_label(table.cell(1, 2), "风力")  # 补充表头
        self.set_table_label(table.cell(1, 3), "气温")  # 补充表头
        # 填充内容
        for i in range(5):
            col = table.columns[i]
            for j in range(7):
                if i == 0:  # 第一列 日期
                    text = f"{weather_info['date'][j]}日"
                elif i == 1:  # 第二列 天气
                    with conn.cursor() as cur:
                        wep_max = int(weather_info["wep_max"][j])
                        wep_min = int(weather_info["wep_min"][j])
                        sql = f"SELECT label FROM weather_code WHERE code in ('{wep_max}', '{wep_min}')"
                        cur.execute(sql)
                        wep_labels = [_[0]for _ in cur.fetchall()]
                    text = wep_labels[0] if len(wep_labels) == 1 else "转".join(wep_labels)
                elif i == 2:  # 第三列 风力
                    win_s_min = weather_info["win_s_min"][j]
                    win_s_max = weather_info["win_s_max"][j]
                    win_l_min, win_l_max = self.get_win_level(win_s_min, win_s_max)
                    text = f"{win_l_min}-{win_l_max}级" if win_l_min != win_l_max else f"{win_l_min}级"
                elif i == 3:  # 第四列 气温
                    tem_min = weather_info["tem_min"][j]
                    tem_max = weather_info["tem_max"][j]
                    text = f"{tem_min}～{tem_max}℃"
                else:  # 第五列 风向
                    win_d_max = weather_info["win_d_max"][j]
                    win_d_min = weather_info["win_d_min"][j]
                    win_d = (win_d_min + win_d_max)/2
                    text = self.get_win_direction(win_d)
                self.set_table_cell(col.cells[j + 2], text)

        paragraph = document.add_paragraph(f"    三、{self.region['name']}未来三天逐3小时精细预报", style="New Heading P")
        paragraph.paragraph_format.space_before = Pt(16)  # 段前1行

        document.save('./{}.docx'.format(self.uuid))  # universally unique identifier 在当前文件夹下创立一个文件


content_head = u"""摘要：春播期间，主要农区呈现“气温起伏较大，总体偏高，透雨出现及时，大部地区墒情适宜”的特点，对主要作物播种、出苗及苗期生长总体有利。"""

content_p1 = u"""    自3月中旬开始春播以来，我区平均气温东部偏南及西部大部地区为10～14℃，其余地区为1～9℃（图1），与历年同期相比，大部地区接近常年或偏高0.6～2.2℃（图2）；东部偏南地区、东部偏东地区及西部个别地区累计降水量为51～135mm，其余地区为7～49mm（图3），中西部偏南地区及东部个别地区较常年偏少1成～6成，其余大部地区偏多1成～2倍（图4）。
3月中旬至5月下旬大于10℃有效积温，东部偏南及西部大部地区为202～472℃·d，其余地区为34～199℃·d（图5）；气温稳定通过0℃初日，西部及东部偏南地区出现在3月上中旬，中东部偏北地区出现在4月上旬，东北部出现在4月中旬，东部偏南、西部大部农区及东北部个别农区偏早2～17天；气温稳定通过10℃初日,西部及东部偏南大部地区出现在4月中旬，其余大部农区为5月中旬，东部偏南、中西部个别农区较常年偏早2～14天。
春播期，主要农区热量总体充足，东部大部墒情适宜，利于主要作物播种及幼苗生长。4月上中旬及5月上中旬的阶段性低温，对东部区春小麦及春玉米的播种及出苗不利；中西部部分农区持续缺墒，对旱作区作物播种和幼苗生长有不利影响。
"""

content_p2 = u"""    我区春小麦于3月中旬开始零星播种，4月上旬开始大面积播种，播种进度较快；河套灌区4月上旬末完成播种；中部区5月上中旬完成播种。春玉米于4月中下旬自西向东陆续开展，5月下旬播种接近尾声。总体来看，自春播以来，中西部部分农区墒情持续偏差，水分条件不足对小麦出苗、叶片生长、分蘖形成和玉米苗期生长造成不利影响；东部区前期气温持续偏低，气象条件影响春耕整地，但后期墒情适宜，有利于加快播种进度，目前苗情较好。
    1.3月中下旬大部农区热量充足，河套灌区墒情总体适宜，麦播开展顺利
3月中下旬，我区气温偏高到特高，大部地区无明显降水，春小麦播种自西向东陆续展开。西部地区回温较快，主要农区平均气温达6～11℃，尤其河套灌区热量充足，土壤水分条件较历年相比较好，利于春小麦大面积播种，东部偏南地区墒情偏差，对春小麦播种略有不利；东北部积雪深度较厚，土壤底墒充足，春播生产条件适宜。
    2.4月上中旬主要农区气温偏低，第一场透雨及时，墒情适宜，总体利于春播
4月上中旬气温偏低、降水偏多，4月1日至3日，全区迎来入春以来第一场大范围降水天气过程，除东北部地区外，大部地区累计降水量为10～50mm，中部及东部偏南大部地区前期偏差墒情得到明显改善，为春播和已播作物出苗创造了有利的土壤水分条件；河套灌区潮塌加重，不利于麦播收尾工作和已播小麦出苗，影响出苗速度和出苗率。
    3.4月下旬大部地区温高光足，气象条件利于春播
4月下旬多晴好天气，温高光足，利于玉米、马铃薯、大豆等作物大面积播种，温高雨少的天气对前期墒情偏差的农区保墒略有不利影响，大部地区光、热、水条件匹配较好，对春播生产较为有利。
    4.5月上中旬降水分布不均，气温略偏低，对春播影响不大 
中东部大部地区较明显的降温、降水天气过程，利于前期墒情偏差的农区增墒保墒和已播作物的生长，但东部偏南个别地区受低温阴雨天气影响，热量略有不足，不利于玉米、马铃薯和大豆等作物的播种作业，对已播小麦的出苗及幼苗生长影响不大。 
5.5月下旬气温逐步回升，中西部部分农区持续缺墒，对旱作区作物播种和幼苗生长有不利影响
大部地区温高光足，但中西部大部基本无降水，对旱作农区小麦分蘖、拔节及玉米苗期生长略有不利影响，苗情多为二类苗。
从5月29日农区土壤墒情监测结果来看，农区一类、二类、三类墒情面积分别为1.1、3.7、4.7万平方公里，分别占农区总面积的11.3%、38.9%、49.8%（图6）。三类墒情面积比去年同期增加1个百分点，比历年同期减少6个百分点（表）。农区墒情接近去年和历年同期。墒情偏差的农区主要分布在乌兰察布市、包头市、呼和浩特市、鄂尔多斯市，对上述地区已播作物的顺利出苗和叶片生长略有不利影响。
"""

content_p3 = u"""    目前，我区大部农区春播顺利结束，呼伦贝尔市、锡林郭勒盟及乌兰察布市各类作物进入春播收尾阶段。农业气象观测显示，目前，我区春小麦河套灌区、东部偏南大部农区处于孕穗期，阴山北麓东段处于三叶至分蘖期，苗情以二类苗为主（图7）；春玉米东部偏南部分地区处于三叶期，其余大部农区处于出苗到幼苗生长阶段，苗情以一、二类苗为主（图8）。大部农区春小麦发育期接近常年或偏早2～6天；春玉米大部农区接近常年或偏早4～5天，受前期土壤墒情偏差影响东部偏南个别农区发育期偏晚4～8天。"""

content_p4 = u"""    目前我区春播已基本结束，建议各地密切关注天气变化，尚未完成播种的地区抓住晴好天气，抢播抢种，尽快完成计划播种任务。据自治区气象台预报， 6月上旬中西部大部农区无明显降水过程，大部农区温高雨少，土壤墒情将继续下滑，建议旱作农区结合当前土壤墒情及天气形势的发展变化，关注墒情变化及对作物苗期生长的影响，做好抗旱预防和准备工作。 """

if __name__ == '__main__':
    json_path = sys.argv[1:][0]
    if not os.path.exists(json_path):
        raise ValueError("Not Found Json Config")
    with open(json_path, "rb") as fp:
        # fp.read()  # string
        argv = json.loads(fp.read())  # dict
    doc = Doc141(**argv)
    if doc.region["level"] == 2:
        doc.make_doc_for_country()
    else:
        doc.make_doc()
