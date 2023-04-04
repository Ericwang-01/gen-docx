# -*- coding: utf-8 -*-
from __future__ import print_function
import os
import json
import math
import psycopg2
import requests
from datetime import datetime
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from meteor_data.daily_data import DailyDataModel


class BaseProduct(object):
    MEADOW = {1: u"上旬", 2: u"中旬", 3: u"下旬"}
    SEASON = {"春": 1, "夏": 2, "秋": 3, "冬": 4}
    WEP = {0: "晴", 1: "多云", 2: "阴", 3: "阵雨", 4: "雷阵雨", 5: "雷阵雨伴有冰雹", 6: "雨夹雪", 7: "小雨", 8: "中雨", 9: "大雨"}
    TIF_DIR = {1: "common_region_data", 2: "common_product_data"}

    def __init__(self, *args, **kwargs):
        self.content_head = ""
        self.content_temperature = ""
        self.content_rainfall = ""
        self.content_impact = ""
        self.content_sunlight = ""
        self.content_soil_moisture = ""
        self.data = None
        self.root_path = kwargs.get("root_path")
        self.owner = kwargs.get("owner")
        self.signer = kwargs.get("signer")
        self.main_class = kwargs.get("analysis")
        self.chief = kwargs.get("reviewer")
        self.unit = kwargs.get("unit")
        self.unit_pic = kwargs.get("unit_pic")
        self.label = kwargs.get("label")  # 标题
        self.date = kwargs.get("date_now").split("-")  # 当前时间列表
        self.period = kwargs.get("period")
        self.uuid = kwargs.get("my_uuid")
        self.issue = kwargs.get("issue")
        self.issue_total = kwargs.get("issue_total")
        self.dsn = kwargs.get("dsn")
        self.dsn_sys = kwargs.get("dsn_sys")

        self.station_dict = kwargs.get("station_dict")
        self.station_location = kwargs.get("station_location")
        self.region = kwargs.get("region")

    @staticmethod
    def load_style(document):
        """
        Loads Global Word Document Style
        :return:
        """
        # 设置页面边距
        document.sections[0].top_margin = Cm(2.3)
        document.sections[0].bottom_margin = Cm(2.4)
        document.sections[0].left_margin = Cm(2.5)
        document.sections[0].right_margin = Cm(2.5)
        # WORD 报头标题 红字
        document.styles.add_style("Word Header", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["Word Header"].font.name = u"华文新魏"
        document.styles["Word Header"].element.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')
        document.styles["Word Header"].font.bold = True
        document.styles["Word Header"].font.size = Pt(36)  # 小初
        document.styles["Word Header"].font.color.rgb = RGBColor(255, 0, 0)  # 红色
        document.styles["Word Header"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 段 标题
        document.styles.add_style('Heading P', WD_STYLE_TYPE.PARAGRAPH)
        document.styles['Heading P'].font.name = u'宋体'
        document.styles['Heading P'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Heading P'].font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        document.styles['Heading P'].font.size = Pt(14)  # 四号
        document.styles['Heading P'].font.bold = True
        # 期刊号
        document.styles.add_style("Issue", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["Issue"].font.name = u'楷体_GB2312'
        document.styles["Issue"].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体_GB2312')
        document.styles["Issue"].font.size = Pt(16)  # 三号
        document.styles["Issue"].font.bold = True
        document.styles["Issue"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 默认
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        document.styles['Normal'].font.size = Pt(12)  # 小四
        # 大标题
        document.styles['Body Text 3'].font.name = u'华文新魏'
        document.styles['Body Text 3'].font.size = Pt(36)  # 小初
        document.styles['Body Text 3'].font.bold = True
        document.styles['Body Text 3'].font.color.rgb = RGBColor(255, 0, 0)  # 红色
        document.styles['Body Text 3'].element.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')
        # 报文题目
        document.styles['Body Text 2'].font.name = u'楷体_GB2312'
        document.styles['Body Text 2'].font.size = Pt(16)  # 三号
        document.styles['Body Text 2'].font.bold = True
        document.styles['Body Text 2'].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体_GB2312')
        # 摘要
        document.styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
        document.styles['Abstract'].font.name = u'仿宋_GB2312'
        document.styles['Abstract'].element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        document.styles['Abstract'].font.size = Pt(14)  # 四号
        document.styles['Abstract'].font.bold = True
        document.styles['Abstract'].font.color.rgb = RGBColor(128, 0, 0)  # 暗红色
        document.styles['Abstract'].font.highlight_color = WD_COLOR_INDEX.GRAY_25  # 灰底
        # 签发人行
        document.styles.add_style("SignerLine", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["SignerLine"].font.name = u'楷体_GB2312'
        document.styles["SignerLine"].font.size = Pt(16)  # 三号
        document.styles["SignerLine"].font.bold = True
        document.styles["SignerLine"].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体_GB2312')
        document.styles["SignerLine"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.styles["SignerLine"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 单倍行距
        document.styles["SignerLine"].paragraph_format.space_after = Pt(0)  # 段落间距

        document.styles.add_style("label", WD_STYLE_TYPE.PARAGRAPH)
        document.styles['label'].font.name = u'宋体'
        document.styles['label'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['label'].font.bold = True
        document.styles['label'].font.size = Pt(16)
        # 图片说明
        document.styles.add_style("Pic Desc", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["Pic Desc"].font.name = u'黑体'
        document.styles["Pic Desc"].element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.styles["Pic Desc"].font.size = Pt(10.5)
        document.styles["Pic Desc"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 图片说明
        document.styles.add_style("pic_desc", WD_STYLE_TYPE.PARAGRAPH)
        document.styles['pic_desc'].font.name = u'黑体'
        document.styles['pic_desc'].element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.styles['pic_desc'].font.size = Pt(10.5)
        # 表头
        document.styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["Table Header"].font.name = u'黑体'
        document.styles["Table Header"].element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.styles["Table Header"].font.size = Pt(14)  # 四号
        document.styles["Table Header"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.styles["Table Header"].font.bold = True
        # 内部资料行
        document.styles.add_style("reference", WD_STYLE_TYPE.PARAGRAPH)
        document.styles["reference"].font.name = u'宋体'
        document.styles["reference"].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles["reference"].font.size = Pt(14)  # 四号
        document.styles["reference"].font.bold = True
        document.styles["reference"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        document.styles["reference"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    @staticmethod
    def get_legend_data(legend_id, conn):
        """
        查询图例
        @param legend_id:
        @param conn:
        @return:
        """
        sql = """SELECT id,label,data,unit FROM legend WHERE id={0}""".format(legend_id)
        with conn.cursor() as cur:
            cur.execute(sql)
            # labels = [str(row[0]) for row in cur.description]
            rows = cur.fetchall()
        return rows

    @staticmethod
    def insert_top_pic(document, width, height):
        """
        插入报头图片
        :param height: 图片高度
        :param width: 图片宽度
        :param document:文档对象
        :return:
        """
        new = document.add_paragraph("")
        run = new.add_run()
        run.add_picture(u"./head_title.jpg", width=Cm(width), height=Cm(height))
        new.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def insert_red_line_pic(document, width, height):
        """
        插入报头图片
        :param height: 图片高度
        :param width: 图片宽度
        :param document:文档对象
        :return:
        """
        new = document.add_paragraph("")
        run = new.add_run()
        run.add_picture(u"./red_line.jpg", width=Cm(width), height=Cm(height))
        new.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    @staticmethod
    def insert_pic(document, path, width, height):
        """
        插入图片, 可自定义宽高
        :param document: 文档对象
        :param path: 文件路径
        :param width: 宽度
        :param height: 高度
        :return:
        """
        new = document.add_paragraph("")
        run = new.add_run()
        if not os.path.exists(path.encode("utf-8")):
            run.add_picture(u"./NaN.png", width=Cm(width), height=Cm(height))
        else:
            try:
                run.add_picture(path, width=Cm(width), height=Cm(height))
            except Exception:
                run.add_picture(u"./NaN.png", width=Cm(width), height=Cm(height))
        new.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def insert_pic_two(document, path1, path2, width, height):
        """
        插入并行两张图片, 可自定义宽高
        :param path1: 图片路径1
        :param path2: 图片路径2
        :param document: 文档对象
        :param width: 宽度
        :param height: 高度
        :return:
        """
        new = document.add_paragraph("")
        run = new.add_run()
        if not os.path.exists(path1.encode("utf-8")):
            run.add_picture(u"./NaN.png", width=Cm(width), height=Cm(height))
        else:
            run.add_picture(path1, width=Cm(width), height=Cm(height))
        if not os.path.exists(path2.encode("utf-8")):
            run.add_picture(u"./NaN.png", width=Cm(width), height=Cm(height))
        else:
            run.add_picture(path2, width=Cm(width), height=Cm(height))
        new.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def match_meadow(date):
        """
        当前日期判断 旬
        :return:
        """
        year = int(date[:4])
        month = int(date[5:7])
        day = date[8:10]
        if int(day[0]) == 0:
            return u"{}年{}月上旬".format(year, month)
        elif int(day[0]) == 1:
            return u"{}年{}月中旬".format(year, month)
        elif int(day[0]) in [2, 3]:
            return u"{}年{}月下旬".format(year, month)

    @staticmethod
    def match_tem_status(v):
        """
        匹配 气温距平 要素值
        :param v: 距平值
        :return:
        """
        if -40 < v <= -2:
            return u"异常偏低"
        elif -2 < v < -1:
            return u"偏低"
        elif -1 <= v <= 1:
            return u"接近常年"
        elif 1 < v <= 2:
            return u"偏高"
        elif 2 < v < 20:
            return u"异常偏高"

    @staticmethod
    def match_tem_anomaly(dim, v):
        """
        匹配 气温距平 要素值
        :param dim: 粒度 1时 2日 3候 4旬 5月 6季 7年
        :param v: 距平值
        :return:
        """
        if dim == 4:
            if -40 < v <= -3:
                return u"气温较常年异常偏低"
            elif -3 < v < -1:
                return u"气温较常年偏低"
            elif -1 <= v <= 1:
                return u"气温接近常年"
            elif 1 < v <= 3:
                return u"气温较常年偏高"
            elif 3 < v < 20:
                return u"气温较常年异常偏高"
        if dim in [5, 6, 7]:
            if -40 < v <= -2:
                return u"气温较常年异常偏低"
            elif -2 < v < -1:
                return u"气温较常年偏低"
            elif -1 <= v <= 1:
                return u"气温接近常年"
            elif 1 < v <= 2:
                return u"气温较常年偏高"
            elif 2 < v < 20:
                return u"气温较常年异常偏高"

    @staticmethod
    def match_pre_r_status(v, r_v):
        """
        匹配  常年降水量 要素值
        :param v: 常年降水量
        :param r_v: 降水相对距平
        :return:
        """
        if 10 < v < 1000:
            if -100 < r_v <= -70:
                return u"异常偏少"
            elif -70 < r_v <= -30:
                return u"偏少"
            elif -30 < r_v < -10:
                return u"略偏少"
            elif -10 <= r_v <= 10:
                return u"接近常年"
            elif 10 < r_v <= 30:
                return u"略偏多"
            elif 30 < r_v <= 70:
                return u"偏多"
            elif 70 < r_v < 2000:
                return u"异常偏多"
        elif 25 < v < 3000:
            if -100 < r_v <= -50:
                return u"异常偏少"
            elif -50 < r_v <= -30:
                return u"偏少"
            elif -30 < r_v < -10:
                return u"略偏少"
            elif -10 <= r_v <= 10:
                return u"接近常年"
            elif 10 < r_v <= 30:
                return u"略偏多"
            elif 30 < r_v <= 50:
                return u"偏多"
            elif 50 < r_v < 2000:
                return u"异常偏多"
        else:
            return u""

    @staticmethod
    def match_pre_r_anomaly(dim, v, r_v):
        """
        匹配  常年降水量 要素值
        :param dim: 粒度 1时 2日 3候 4旬 5月 6季 7年
        :param v: 常年降水量
        :param r_v: 降水相对距平
        :return:
        """
        if dim == 4:
            if 10 < v < 1000:
                if -100 < r_v <= -70:
                    return u"降水较常年异常偏少"
                elif -70 < r_v <= -30:
                    return u"降水较常年偏少"
                elif -30 < r_v < -10:
                    return u"降水略偏少"
                elif -10 <= r_v <= 10:
                    return u"降水接近常年"
                elif 10 < r_v <= 30:
                    return u"降水略偏多"
                elif 30 < r_v <= 70:
                    return u"降水较常年偏多"
                elif 70 < r_v < 2000:
                    return u"降水异常偏多"
            else:
                return ""
        if dim in [5, 6, 7]:
            if 25 < v < 3000:
                if -100 < r_v <= -50:
                    return u"降水较常年异常偏少"
                elif -50 < r_v <= -30:
                    return u"降水较常年偏少"
                elif -30 < r_v < -10:
                    return u"降水略偏少"
                elif -10 <= r_v <= 10:
                    return u"降水接近常年"
                elif 10 < r_v <= 30:
                    return u"降水略偏多"
                elif 30 < r_v <= 50:
                    return u"降水较常年偏多"
                elif 50 < r_v < 2000:
                    return u"降水异常偏多"
            else:
                return u""

    @staticmethod
    def match_sun_hour(values):
        """
        日照的百分率  要素值
        :param values: 日照的百分率
        :return:
        """
        if 0 <= values < 20:
            return u"略有不足"
        elif 20 <= values < 60:
            return u"良好"
        elif 60 <= values <= 100:
            return u"充足"
        else:
            return u""

    @staticmethod
    def match_frost_damage(day_tmp_min, g_tmp_min):
        """
        冻害 要素值
        :param day_tmp_min: 日最低气温
        :param g_tmp_min: 最低地表温度
        :return:
        """
        if 0 <= day_tmp_min <= 2:
            if -2 <= g_tmp_min <= 0:
                return u"轻度冻害"
        elif -2 <= day_tmp_min < 0:
            if -4 <= g_tmp_min <= -2:
                return u"中度冻害"
        elif -50 <= day_tmp_min < -2:
            if -50 <= g_tmp_min <= -4:
                return u"中度冻害"

    @staticmethod
    def match_spring_pre(v):
        """
        判断 春播 降水的状态
        :param v:
        :return:
        """
        if -20 > v:
            return u"偏少"
        elif -20 < v < 50:
            return u"适宜"
        elif v > 50:
            return u"偏多"

    @staticmethod
    def compute_sun_percent(date, lat, sun_hour):
        """
        计算日照 百分率
        :param sun_hour: 日照时数
        :param date: 日期
        :param lat: 维度
        :return:
        """
        if isinstance(date, str):
            date = datetime.strptime(date, "%Y-%m-%d")
        total_days = date.timetuple().tm_yday
        # sun_ground_distance = 1 + 0.033 * math.cos(2 * math.pi * total_days / 365)
        sun_lat = 0.409 * math.sin(2 * math.pi * total_days / 365 - 1.39)
        sunset_view = math.acos(-math.tan(math.pi / 180 * lat) * math.tan(sun_lat))
        sun_percent = 24 / math.pi * sunset_view
        return round(sun_hour / sun_percent * 100, 1)

    @staticmethod
    def format_interval(v1, v2):
        if v1 == v2:
            return str(v1)
        return "{}～{}".format(v1, v2)

    @staticmethod
    def predict_ssp_level(v):
        if v > 70:
            return "日照充足"
        elif v < 30:
            return "日照偏少"
        else:
            return "日照良好"

    @staticmethod
    def match_season(date):
        """
        输入日期判断 季度  与message接口不同
        :param date: str类型日期
        :return:
        """
        year = int(date[:4])
        month = int(date[5:7])
        if month == 12:
            return "{}年冬季".format(year)
        elif month in [3, 4, 5]:
            return "{}年春季".format(year)
        elif month in [6, 7, 8]:
            return "{}年夏季".format(year)
        elif month in [9, 10, 11]:
            return "{}年秋季".format(year)
        elif month in [1, 2]:
            return "{}年冬季".format(year - 1)

    @staticmethod
    def match_period(st, et):
        """
        Match StartTime EndTime Format String Type Period
        @param st:
        @param et:
        @return:
        """
        if isinstance(st, str):
            st = datetime.strptime(st, "%Y-%m-%d")
        if isinstance(et, str):
            et = datetime.strptime(et, "%Y-%m-%d")
        if st.month == et.month:
            if st.day == et.day:
                return f"{et.month}月{st.day}日"
            return f"{et.month}月{st.day}日~{et.day}日"
        return f"{st.month}月{st.day}日~{et.month}月{et.day}日"

    def draw_common_png(self, *args, **kwargs):
        """
        绘制八要素png图片
        @param args:
        @param kwargs:  station_type --- 1全部站2国家站
                        conn --- 数据库连接
                        legend_id --- 查询图例的ID
                        save_path --- 保存 png图片的路径,填空字符串返回文件流
                        tm --- 时间粒度1-7
                        ele_name --- 要素名称
                        title --- 标题
        @return:
        """
        s_type = kwargs.setdefault("station_type", 2)
        conn = kwargs.get("conn")
        legend_id = kwargs.get("legend_id")
        save_path = kwargs.get("save_path")
        tm = kwargs.get("tm")
        ele_name = kwargs.get("ele_name")
        title = kwargs.get("title")
        tif_date_ten_month = self.period[0][:7].replace("-", '')
        tif_date_season_year = self.period[0][:4]

        if tm == 4:
            day = int(self.period[0][-2:])
            if day == 1:
                xun = "上旬"
                tm_detail = 1
            elif day == 11:
                xun = "中旬"
                tm_detail = 2
            else:
                xun = "下旬"
                tm_detail = 3
            date_title = u"{}年{}月{}".format(self.period[0][:4], self.period[0][5:7], xun)
            tif_url = f"{self.TIF_DIR[s_type]}/{tm}/{tif_date_ten_month}/{tm_detail}/{ele_name}.tif"
        elif tm == 5:
            date_title = u"{}年{}月".format(self.period[0][:4], self.period[0][5:7])
            tif_url = f"{self.TIF_DIR[s_type]}/{tm}/{tif_date_ten_month}/{1}/{ele_name}.tif"
        elif tm == 6:
            date_title = self.match_season(self.period[0])
            tm_detail = self.SEASON[date_title[-2]]
            tif_url = f"{self.TIF_DIR[s_type]}/{tm}/{tif_date_season_year}/{tm_detail}/{ele_name}.tif"
        elif tm == 7:
            date_title = u"{}年".format(self.period[0][:4])
            tif_url = f"{self.TIF_DIR[s_type]}/{tm}/{tif_date_season_year}/{1}/{ele_name}.tif"
        else:
            raise KeyError("Not Support Current Size Of Time")
        f_url = os.path.join(self.root_path, tif_url)
        print(f_url)

        return self.curl_node_to_pic(f_url, save_path, legend_id, conn, title, date_title)

    def curl_node_to_pic(self, tif_url, save_path, legend_id, conn, title, date_title):
        """
        Curl Node Server Response Save To PNG
        @param tif_url:
        @param save_path:
        @param legend_id:
        @param conn:
        @param title:
        @param date_title:
        @return:
        """
        if not tif_url.startswith("http") and not os.path.exists(tif_url):
            return
        rows = self.get_legend_data(legend_id, conn)
        legend_data = rows[0][2]  # 图例数据
        unit = rows[0][3]  # 图例单位
        legend = json.loads(legend_data)[0]
        legend["show"] = 1  # 控制图例是否显示
        png_data = {
            "regioncode": self.region["code"],
            "savepath": save_path,
            "tifurl": tif_url,
            "legend": legend,
            "company": {"icon": self.unit_pic, "label": self.unit},
            "title": {"text": title, "subtext": date_title},
            "unit": {"text": unit},
            "projecctionStyle": "geoMercator"  # TODO Wrong Spell projection
        }
        # url = "http://172.18.147.13:9000/api/createGis"
        url = "http://192.168.3.202:9000/api/createGis"
        res = requests.post(url, json=png_data)
        print(res.status_code)
        if res.status_code == 500:
            print(res.content)
            return None
        return res.content

    def format_rainfall(self, v_list, max_station):
        if v_list[0] == v_list[-1]:
            if v_list[0] == 0:
                return "全{}无降水".format(self.region["name"][-1])
            return "全{}为{}mm".format(self.region["name"][-1], v_list[0])
        return "降水量最大值出现在{}，累计降水量为{}mm，其余地区为{}mm（图5）".format(
            max_station,
            v_list[-1],
            self.format_interval(v_list[0], v_list[-2])
        )

    def set_footer(self, document):
        """
        设置页脚内容与红线
        """
        if self.region["level"] == 0:
            text = f"{self.unit}                    分析：{self.main_class}                       审核：{self.chief}"
        elif self.region["level"] == 1:
            blank = " " * (20 - len(self.main_class) - len(self.chief))
            text = f"""报：市委、政府、区局科技预报处、区局应急与减灾处、区局生态与农业气象中心\n市农牧业局、相关科室{blank}分析：{self.main_class}  审核：{self.chief}"""
        else:
            text = f"""报：旗委、政府、旗农牧业局、市局领导、市气象台、业务管理、业务科\n分析：{self.main_class}  审核：{self.chief}"""

        footer = document.sections[0].footer.paragraphs[0]
        # 页脚红线
        footer.add_run().add_picture(u"./thin_red_line.jpg", width=Cm(16.56), height=Cm(0.03))
        # 页脚文本
        footer.add_run(text)

    def set_sign_date(self, document):
        # 2022-03-14 update:字符长度,目前word显示正常,转换pdf仍会串行
        # 统计为一行58个英文字符 - 年月日(6) - 签发：(6) - 保留(4)
        dt = f"{self.date[0]}年{int(self.date[1])}月{int(self.date[2])}日"
        dt_len = len(self.date[0] + str(int(self.date[1])) + str(int(self.date[2])))
        if self.region["level"] == 0:
            blank = " " * (42 - dt_len - len(self.signer) * 2)
            sign_date = document.add_paragraph(f"""{dt}{blank}签发：{self.signer}""", style="SignerLine")
        else:
            blank = " " * (42 - len(self.unit) * 2 - len(self.signer) * 2)
            sign_date = document.add_paragraph(f"""{self.unit}{blank}签发：{self.signer}\n{dt}""", style="SignerLine")

    def set_reference(self, document):
        if self.region["level"] == 0:
            return
        document.add_paragraph("""内部资料\n仅供参考""", style="reference")

        blank = document.add_paragraph("")
        blank.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 单倍行距
        blank.paragraph_format.space_after = Pt(41)  # 段落间

    def set_issue(self, document, content):
        """{}年秋收气象服务专报"""
        if self.region["level"] == 0:
            return document.add_paragraph(f"{content}—第{self.issue}期   （总{self.issue_total}期）", style="Issue")
        return document.add_paragraph(f"第{self.issue}期", style="Issue")

    def gen_base_meteo_data(self):
        conn = psycopg2.connect(self.dsn)
        meteo_obj = DailyDataModel(["tem_avg", "ssh"], self.station_dict.keys(), self.period, conn)
        data = meteo_obj.fetch_product_data()
        data.insert(len(data.columns) - 2, "pre_time_2020_r_anomaly",
                    round(data["pre_time_2020_anomaly"] / data["pre_time_2020_mmut"], 1))
        ssp_map = zip(data["datetime"], data["station_id_c"], data["ssh"])
        ssp = [round(self.compute_sun_percent(dt, self.station_location[s_id][0], ssh), 1) for dt, s_id, ssh in ssp_map]
        data.insert(len(data.columns) - 1, "ssp", ssp)

        return data
