# -*- coding:utf-8 -*-
import os
import json
from datetime import datetime
import psycopg2
import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

matplotlib.use("AGG")

from matplotlib import pyplot as plt
from matplotlib.font_manager import FontProperties
from meteor_data.avg30y_data import Avg30yDataModel

# FONT_PATH = "."
# SAVE_PATH = "static/image"
# DSN = 'host=127.0.0.1 port=5433 dbname=nmnq_db user=nmnq password=nmnq'

FONT_PATH = "/data1/font"
SAVE_PATH = "/data1/CR/product_image"
DSN = 'host=192.168.3.218 dbname=nmnq_db user=nmnq password=nmnq'

dis_translate_dic = {
    '站点号': 'station', '站点': 'station_name',
    '日期': 'date', '年': 'year', '月': 'month', '日': 'day',
    '平均温度(℃)': 'tem_avg', '最高温度(℃)': 'tem_max', '最低温度(℃)': 'tem_min',
    '平均湿度(%)': 'rhu_avg',
    '日照时数(h)': 'ssh', '日照百分率(%)': 'ssp',
    '20-20时降水(mm)': 'pre_time_2020',
    '最大风速(m/s)': 'win_s_max',
}


# 降水日数
def cal_pre_days(pre_list):
    days = 0
    for pre in pre_list:
        if pre > 0:
            days += 1
    return days


# 计算大风日数
def cal_wind_days(wind_list):
    days = 0
    for wind in wind_list:
        if wind == True:
            days += 1
    return days


# 对基本气象数据聚合，生成关键数据
def generate_group_data(df):
    # 获取最高温度所在行信息
    tem_max_df = df[df['tem_max'] == np.max(df['tem_max'])]
    # 获取win_s_max所在行信息
    wind_df = df[df['win_s_max'] == 1].groupby(['year', 'month', 'day']).agg({'win_s_max': 'min'})
    wind_df.reset_index(inplace=True)

    # 聚合
    g_df = df.groupby(['station', 'station_name', 'year', 'month'], sort=False).agg({
        'tem_avg': 'mean',
        'ssh': 'sum',
        'pre_time_2020': ['sum', cal_pre_days],
        'tem_max': 'max',
        'tem_min': 'min',
        'ssp': 'mean',
        'win_s_max': cal_wind_days
    })
    g_df.columns = ['tem_avg', 'ssh', 'pre', 'pre_nums', 'tem_max', 'tem_min', 'ssp', 'win_s_max_nums']
    g_df.reset_index(drop=False, inplace=True)
    g_df = g_df[
        ['station', 'station_name', 'tem_avg', 'ssh', 'pre', 'tem_max', 'tem_min', 'ssp', 'pre_nums', 'win_s_max_nums']]

    return g_df, tem_max_df, wind_df


# 获取同期数据
def get_mmut_data(stations, period):
    tm = 2
    db_connect = psycopg2.connect(DSN)
    a30y_data = Avg30yDataModel(['tem_avg', 'ssh', 'pre_time_2020'], stations, period, db_connect).fetch_data(tm)
    return a30y_data


# 汇总数据
def generate_summarh_day(g_data, mmut_data):
    # 同期数据聚合
    mmut_data = mmut_data.groupby('station').agg({'tem_avg': 'mean', 'ssh': 'sum', 'pre_time_2020': 'sum'})
    mmut_data.columns = ['tem_avg_mmut', 'ssh_mmut', 'pre_mmut']

    # 合并数据
    data = g_data.join(mmut_data)
    data.reset_index(inplace=True)

    # 计算距平
    data['tem_avg_anomaly'] = data['tem_avg'] - data['tem_avg_mmut']
    data['ssh_anomaly'] = data['ssh'] - data['ssh_mmut']
    data['pre_anomaly'] = round((data['pre'] - data['pre_mmut']) / data['pre_mmut'] * 100)
    return data


# 生成首段描述
def generate_description(df, tem_max_df, wind_df, term):
    """
    params:
        df(pd.DataFrame): 站点聚合数据
        term(str): 时期，月或旬
        tem_max_df(pd.DataFrame): 最高气温数据
        wind_df(pd.DataFrame): 大风数据
    return:
        description(str): 首段描述
    """
    des = ''
    tem_avg_min, tem_avg_max = round(np.min(df['tem_avg']), 1), round(np.max(df['tem_avg']), 1)
    des += f'本{term}平均气温{tem_avg_min}~{tem_avg_max}℃'
    his_tem_min, his_tem_max = round(np.min(df['tem_avg_mmut']), 1), round(np.max(df['tem_avg_mmut']), 1)
    des += f'（常年平均{his_tem_min}~{his_tem_max}℃），'
    tem_anomaly_min, tem_anomaly_max = round(np.min(df['tem_avg_anomaly']), 1), round(np.max(df['tem_avg_anomaly']), 1)
    des += f'{term}气温距平{tem_anomaly_min}~{tem_anomaly_max}℃；'
    tem_max = round(np.max(df['tem_max']), 1)
    tem_max_day = ''
    for _, row in tem_max_df.iterrows():
        tem_max_day += f'{row["month"]}月{row["day"]}日的{row["station"]}、'
    des += f'极端最高气温出现在{tem_max_day[:-1]}，为{tem_max}℃；'

    his_pre_max = round(np.max(df['pre_mmut']), 1)
    his_pre_min = round(np.min(df['pre_mmut']), 1)
    his_pre_s = f'（常年平均{his_pre_min}~{his_pre_max}毫米）'
    pre_max = round(np.max(df['pre']), 1)
    pre_min = round(np.min(df['pre']), 1)
    pre_days_min = np.min(df['pre_nums'])
    pre_days_max = np.max(df['pre_nums'])
    if pre_max == 0:
        pre_s = f'{term}内无降水{his_pre_s}；'
    else:
        s1 = f'{term}降水量{pre_min}～{pre_max}毫米{his_pre_s}，'
        s2 = f'{term}降水量距平百分率{np.min(df["pre_anomaly"])}～{np.max(df["pre_anomaly"])}%，'
        pre_days = pre_days_min if pre_days_min == pre_days_max else f'{pre_days_min}~{pre_days_max}'
        s3 = f'{term}降水日数{pre_days}天；'
        pre_s = s1 + s2 + s3
    des += pre_s

    ssh_min, ssh_max = round(np.min(df['ssh']), 1), round(np.max(df['ssh']), 1)
    des += f'{term}日照时数{ssh_min}~{ssh_max}小时 '
    his_ssh_min, his_ssh_max = round(np.min(df['ssh_mmut']), 1), round(np.max(df['ssh_mmut']), 1)
    des += f'（常年平均{his_ssh_min}~{his_ssh_max}小时），'
    ssh_anomaly_min, ssh_anomaly_max = round(np.min(df['ssh_anomaly']), 1), round(np.max(df['ssh_anomaly']), 1)
    des += f'{term}日照距平{ssh_anomaly_min}~{ssh_anomaly_max}小时；'

    wind_days_max = np.max(df['win_s_max_nums'])
    wind_days_min = np.min(df['win_s_max_nums'])
    if wind_days_max == 0:
        wind_s = f'{term}内无7级以上大风。'
    else:
        if term == '月':
            wind_days = ''
            i = 0
            for _, row in wind_df.iterrows():
                if wind_days == '':
                    wind_days += f'{row["month"]}月'
                wind_days += f'{row["day"]}日、'
                i += 1
            wind_s = f'{term}7级以上大风日数{i}天'
            wind_s += f'，出现在{wind_days[:-1]}。'
        else:
            wind_days = wind_days_min if wind_days_min == wind_days_max else f'{wind_days_min}~{wind_days_max}'
            wind_s = f'{term}7级以上大风日数{wind_days}天。'
    des += wind_s

    return des


# 绘制均温图像
def paint_tem_avg(m_data, mmut_data, label, font):
    # 近期数据
    m_data = m_data.groupby(['year', 'month', 'day', 'date']).agg({'tem_avg': 'mean'})
    m_data.reset_index(inplace=True)
    m_data['date_seq'] = [str(datetime.strptime(date, '%Y-%m-%d').timetuple().tm_yday) for date in m_data['date']]
    m_data.set_index(['date_seq'], inplace=True)

    # 同期数据
    mmut_data = mmut_data.groupby(['date']).agg({'tem_avg': 'mean'})
    mmut_data.reset_index(inplace=True)
    mmut_data.columns = ['date_seq', 'tem_avg_mmut']
    mmut_data.set_index(['date_seq'], inplace=True)
    mmut_data = mmut_data[1:]
    mmut_data = mmut_data[:m_data.shape[0]]

    # 画板设置
    date_list = [str(int(m_data['month'][i])) + '/' + str(int(m_data['day'][i])) for i in
                 range(m_data.shape[0])]  # 获取日期列表
    plt.figure(figsize=(len(date_list) // 2, 5), dpi=100)

    # 设置y轴最值与步长
    tem_y_min = min(min(m_data['tem_avg']), min(mmut_data['tem_avg_mmut']))
    tem_y_max = max(max(m_data['tem_avg']), max(mmut_data['tem_avg_mmut']))
    tem_step = (tem_y_max - tem_y_min) // 3
    if tem_step == 0:
        tem_step = 1
    tem_y_min = tem_y_min // tem_step * tem_step
    tem_y_max = (tem_y_max // tem_step + 1) * tem_step
    plt.yticks(np.arange(tem_y_min, tem_y_max, tem_step))

    # 绘制折线图
    plt.plot(date_list, m_data['tem_avg'], marker='o', markersize=3, label=label, color='#3D6AB2')
    plt.plot(date_list, mmut_data['tem_avg_mmut'], marker='o', markersize=3, label='常年同期', color='#B33A3A')

    # 其他设置
    plt.legend(ncol=2, loc='upper right', frameon=False, prop=font)  # 设置图例
    plt.ylabel('平均气温/℃', fontproperties=font)  # 纵轴名称
    plt.grid(axis='y', alpha=0.5)  # 纵坐标网格
    plt.tick_params(labelsize=8)  # 坐标轴刻度字体大小

    # 保存图像
    img_path = os.path.join(SAVE_PATH, f"{label}平均气温折线图.png")
    plt.savefig(img_path, dpi=200)
    return img_path


# 绘制降水量图像
def paint_pre(s_data, label, font):
    # 设置x轴位置与坐标
    station_list = s_data['station_name']
    x = np.arange(len(station_list))
    _, ax = plt.subplots()
    ax.set_xticks(x)
    ax.set_xticklabels(station_list, fontproperties=font)

    # 设置y轴最值与步长
    pre_y_min = min(min(s_data['pre']), min(s_data['pre_mmut']))
    pre_y_max = max(max(s_data['pre']), max(s_data['pre_mmut']))
    pre_step = (pre_y_max - pre_y_min) // 2
    if pre_step == 0:
        pre_step = 1
    pre_y_min = pre_y_min // pre_step * pre_step
    pre_y_max = (pre_y_max // pre_step + 1) * pre_step
    plt.yticks(np.arange(pre_y_min, pre_y_max, pre_step))

    # 绘制并列柱状图
    width = 0.2
    plt.bar(x - width, s_data['pre'], width=width, label=label, color='#24728C', zorder=10)
    plt.bar(x + width, s_data['pre_mmut'], width=width, label='常年同期', color='#B6D086', zorder=10)

    # 其他设置
    plt.legend(ncol=2, loc='upper right', frameon=False, prop=font, handlelength=0.65)  # 设置图例, handlelength为Label前图标长度
    plt.ylabel('降水量/mm', fontproperties=font)  # 纵轴名称
    plt.grid(axis='y', zorder=0, alpha=0.5)  # 纵坐标网格,图层等级0、透明度0.5
    plt.tick_params(labelsize=8)  # 坐标轴刻度字体大小

    # 保存图像
    img_path = os.path.join(SAVE_PATH, f"{label}降水量柱状图.png")
    plt.savefig(img_path, dpi=200)
    return img_path


# 绘制综合数据图像
def pain_summary(m_data, label, font):
    # 近期数据
    m_data = m_data.groupby(['year', 'month', 'day', 'date']).agg({
        'tem_avg': 'mean',
        'tem_max': 'mean',
        'tem_min': 'mean',
        'ssh': 'mean'
    })
    m_data.reset_index(inplace=True)

    # 画板设置
    date_list = [str(int(m_data['month'][i])) + '/' + str(int(m_data['day'][i])) for i in
                 range(m_data.shape[0])]  # 获取日期列表
    fig = plt.figure(figsize=(len(date_list) // 2, 5), dpi=100)
    # 日照坐标轴
    ax_ssh = fig.add_subplot()
    # 温度坐标轴
    ax_tem = ax_ssh.twinx()

    # 设置温度y轴最值与步长
    tem_y_min = min(m_data['tem_min'])
    tem_y_max = max(m_data['tem_max'])
    tem_step = (tem_y_max - tem_y_min) // 5
    if tem_step == 0:
        tem_step = 1
    tem_y_min = tem_y_min // tem_step * tem_step
    tem_y_max = (tem_y_max // tem_step + 1) * tem_step
    ax_tem.set_yticks(np.arange(tem_y_min, tem_y_max, tem_step))

    # 绘制柱状图
    ax_ssh.bar(date_list, m_data['ssh'], width=0.4, label='日照时数', color='#D08281')
    # 绘制折线图
    ax_tem.plot(date_list, m_data['tem_max'], marker='o', markersize=3, label='最高气温', color='#DE5600')
    ax_tem.plot(date_list, m_data['tem_min'], marker='o', markersize=3, label='最低气温', color='#24728C')
    ax_tem.plot(date_list, m_data['tem_avg'], marker='o', markersize=3, label='平均气温', color='#648529')

    # 日照y轴置于右侧，温度y轴置于左侧
    ax_ssh.yaxis.set_label_position('right')
    ax_ssh.yaxis.tick_right()
    ax_tem.yaxis.set_label_position('left')
    ax_tem.yaxis.tick_left()
    # 设置纵坐标网格
    ax_ssh.grid(axis='y', alpha=0.5)
    # 设置图例
    fig.legend(ncol=4, loc='center right', frameon=False, prop=font, bbox_to_anchor=(1, 0.97),
               bbox_transform=ax_tem.transAxes)
    # 设置纵轴名称
    ax_ssh.set_ylabel('日照时数/h', fontproperties=font)
    ax_tem.set_ylabel('空气温度/℃', fontproperties=font)
    # 设置坐标轴刻度字体大小
    ax_ssh.tick_params(labelsize=8)
    ax_tem.tick_params(labelsize=8)

    # 保存图像
    img_path = os.path.join(SAVE_PATH, f"{label}综合数据图.png")
    plt.savefig(img_path, dpi=200)
    return img_path


# 生成word文档
def generate_word_doc_and_pdf(title, description, img_list):
    doc = Document()
    doc.add_heading(title, 1)  # 添加标题

    p = doc.add_paragraph(description)  # 添加段落
    # 设置段落
    p.style.font.size = Pt(10.5)  # 文本字体大小
    p.paragraph_format.first_line_indent = p.style.font.size * 2  # 首行缩进

    # 添加图片
    for img in img_list:
        pic = doc.add_picture(img)
        scale = pic.width  # 图片原始宽度
        pic.width = Cm(15.09)  # 设置图片宽度与doc等宽
        scale = pic.width / scale  # 计算图片缩放比例
        pic.height = int(pic.height * scale)  # 按照比例设置图片高度

    # 设置对齐方式
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc_name = f'{title}.docx'
    doc.save(doc_name)  # 保存文档
    # os.system(f'libreoffice --headless --convert-to pdf {doc_name}')  # doc转pdf

    return doc_name, f'{title}.pdf'


# 获取报表处理结果
def deal_report_info(m_data, stations, period, term):
    # 生成聚合气象数据
    g_data, tem_max_data, wind_data = generate_group_data(m_data)
    g_data.set_index('station', inplace=True)

    # 获取同期数据
    mmut_data = get_mmut_data(stations, period)

    # 汇总数据
    s_data = generate_summarh_day(g_data, mmut_data)

    # 生成首段描述
    description = generate_description(
        df=s_data,
        tem_max_df=tem_max_data,
        wind_df=wind_data,
        term=term
    )

    # 图像图例名
    label = '%s年%s月' % (period[0][:4], period[0][5:7] if int(period[0][5:7]) > 9 else period[0][6:7])
    if term == '旬':
        i10days_dict = {1: '上', 2: '中', 3: '下'}
        label += '%s旬' % (i10days_dict[int((int(period[0][-2:]) + 9) / 10) - int(int(period[0][-2:]) / 31)])
    # 字体设置
    font = FontProperties(fname=os.path.join(FONT_PATH, 'SimHei.ttf'), size=9)
    # 绘制图像
    tem_img = paint_tem_avg(m_data, mmut_data, label, font)
    pre_img = paint_pre(s_data, label, font)
    sum_img = pain_summary(m_data, label, font)
    img_list = [tem_img, pre_img, sum_img]

    return description, img_list


# task_param调用
def run():
    with open('task_param.json', 'r') as f:
        argv = json.loads(f.read())

    # 读取参数
    title = argv.get('label')  # 文件名
    stations = argv.get('stations')  # 站点号列表
    time = argv.get('time')  # 时间参数
    period = time.get('period')  # 时间区间
    term = time.get('groupby')  # 周期：旬、月
    m_data = pd.DataFrame(argv.get('report_info'))
    re_columns = {i: dis_translate_dic[i] for i in m_data.columns}
    m_data.rename(columns=re_columns, inplace=True)
    m_data['station'] = m_data['station'].astype(str)

    # 按粒度调整period时间样式
    if term == 4:
        term = '旬'
        period = [period[0][0], period[1][1]]
    elif term == 5:
        term = '月'
        last_day_map = {1: "-31", 2: "-28", 3: "-31", 5: "-31", 7: "-31", 8: "-31", 10: "-31", 12: "-31", }
        if int(period[1][-2:]) in last_day_map.keys():
            last_day = last_day_map[int(period[1][-2:])]
        else:
            last_day = "-30"
        period = [period[0] + "-01", period[1] + last_day]

    # 生成文件时间前缀；description，天气描述；img_list，图片列表
    description, img_list = deal_report_info(m_data, stations, period, term)

    # 生成报告文档 doc,pdf
    doc_name, pdf_name = generate_word_doc_and_pdf(title, description, img_list)

    return doc_name, pdf_name


if __name__ == '__main__':
    run()
